"""
Professional Document Generation Module
Creates ATS-compliant DOCX and PDF documents from AI agent outputs
"""
import os
import json
import tempfile
from datetime import datetime
from typing import Dict, List, Optional
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors

import markdown


class DocumentGenerator:
    """Professional document generator for career optimization outputs"""
    
    def __init__(self, output_dir: str = "/app/output"):
        self.output_dir = Path(output_dir)
        self.temp_dir = Path("/tmp/documents")
        self.temp_dir.mkdir(exist_ok=True)
    
    def generate_resume_docx(self, resume_data: str, filename: str = None) -> str:
        """Generate ATS-compliant resume in DOCX format"""
        if filename is None:
            filename = f"optimized_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = self.temp_dir / filename
        
        # Create new document
        doc = Document()
        
        # Set document margins (ATS-friendly)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Add styles for ATS compliance
        self._add_resume_styles(doc)
        
        # Parse markdown resume content
        lines = resume_data.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                # Main heading (Name)
                name = line[2:].strip()
                heading = doc.add_paragraph(name)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading.style = doc.styles['Heading 1']
                
            elif line.startswith('## '):
                # Section headings
                section_title = line[3:].strip()
                section_para = doc.add_paragraph(section_title)
                section_para.style = doc.styles['Heading 2']
                current_section = section_title.lower()
                
            elif line.startswith('### '):
                # Sub-section headings
                subsection_title = line[4:].strip()
                subsection_para = doc.add_paragraph(subsection_title)
                subsection_para.style = doc.styles['Heading 3']
                
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet points
                bullet_text = line[2:].strip()
                bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
                
            elif line.startswith('**') and line.endswith('**'):
                # Bold text (job titles, company names)
                bold_text = line[2:-2].strip()
                para = doc.add_paragraph()
                run = para.add_run(bold_text)
                run.bold = True
                
            else:
                # Regular text
                if line:
                    para = doc.add_paragraph(line)
                    para.style = 'Normal'
        
        # Save document
        doc.save(str(filepath))
        return str(filepath)
    
    def generate_cover_letter_docx(self, cover_letter_data: Dict, filename: str = None) -> str:
        """Generate professional cover letter in DOCX format"""
        if filename is None:
            filename = f"cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = self.temp_dir / filename
        
        # Create new document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add styles
        self._add_letter_styles(doc)
        
        # Header with date
        date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add spacing
        doc.add_paragraph()
        
        # Cover letter content
        if isinstance(cover_letter_data, dict):
            content = cover_letter_data.get('cover_letter_content', '')
        else:
            content = str(cover_letter_data)
        
        # Parse content and add to document
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style = 'Normal'
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Professional closing
        doc.add_paragraph()
        closing = doc.add_paragraph("Sincerely,")
        closing.style = 'Normal'
        
        # Add signature space
        for _ in range(3):
            doc.add_paragraph()
        
        signature = doc.add_paragraph("[Your Name]")
        signature.style = 'Normal'
        
        # Save document
        doc.save(str(filepath))
        return str(filepath)
    
    def generate_linkedin_profile_docx(self, linkedin_data: Dict, filename: str = None) -> str:
        """Generate LinkedIn optimization guide in DOCX format"""
        if filename is None:
            filename = f"linkedin_optimization_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = self.temp_dir / filename
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph("LinkedIn Profile Optimization Guide")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = doc.styles['Heading 1']
        
        # Add sections based on LinkedIn data
        sections = [
            ("Optimized Headline", linkedin_data.get('optimized_headline', '')),
            ("Professional Summary", linkedin_data.get('professional_summary', '')),
            ("Skills to Add", '\n'.join(f"• {skill}" for skill in linkedin_data.get('skills_to_add', []))),
            ("Experience Enhancements", self._format_experience_enhancements(linkedin_data.get('experience_enhancements', []))),
            ("Keyword Optimization", '\n'.join(f"• {keyword}" for keyword in linkedin_data.get('keyword_optimization', []))),
            ("Networking Strategy", '\n'.join(f"• {strategy}" for strategy in linkedin_data.get('networking_strategy', []))),
            ("Content Suggestions", '\n'.join(f"• {suggestion}" for suggestion in linkedin_data.get('content_suggestions', [])))
        ]
        
        for section_title, section_content in sections:
            if section_content:
                # Add section heading
                heading = doc.add_paragraph(section_title)
                heading.style = doc.styles['Heading 2']
                
                # Add section content
                content_para = doc.add_paragraph(section_content)
                content_para.style = 'Normal'
                
                # Add spacing
                doc.add_paragraph()
        
        # Save document
        doc.save(str(filepath))
        return str(filepath)
    
    def generate_interview_prep_docx(self, interview_data: Dict, filename: str = None) -> str:
        """Generate interview preparation guide in DOCX format"""
        if filename is None:
            filename = f"interview_preparation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = self.temp_dir / filename
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph("Interview Preparation Guide")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = doc.styles['Heading 1']
        
        # Add sections
        sections = [
            ("Behavioral Questions", interview_data.get('behavioral_questions', [])),
            ("Technical Questions", interview_data.get('technical_questions', [])),
            ("Company-Specific Questions", interview_data.get('company_specific_questions', [])),
            ("Questions to Ask the Interviewer", interview_data.get('questions_to_ask', [])),
            ("Preparation Tips", interview_data.get('preparation_tips', [])),
            ("Potential Challenges", interview_data.get('potential_challenges', []))
        ]
        
        for section_title, section_data in sections:
            if section_data:
                # Add section heading
                heading = doc.add_paragraph(section_title)
                heading.style = doc.styles['Heading 2']
                
                # Add questions/content
                for i, item in enumerate(section_data, 1):
                    if isinstance(item, dict):
                        # Question with guidance
                        question_text = item.get('question', item.get('tip', item.get('challenge', str(item))))
                        guidance = item.get('guidance', item.get('answer_framework', ''))
                        
                        q_para = doc.add_paragraph(f"{i}. {question_text}")
                        q_para.style = 'List Number'
                        
                        if guidance:
                            g_para = doc.add_paragraph(f"   Guidance: {guidance}")
                            g_para.style = 'Normal'
                    else:
                        # Simple list item
                        item_para = doc.add_paragraph(f"{i}. {str(item)}")
                        item_para.style = 'List Number'
                
                # Add spacing
                doc.add_paragraph()
        
        # Save document
        doc.save(str(filepath))
        return str(filepath)
    
    def generate_pdf_from_docx(self, docx_path: str) -> str:
        """Convert DOCX to PDF (simplified version)"""
        # For now, return the DOCX path as PDF generation requires additional dependencies
        # In a production environment, you would use python-docx2pdf or similar
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # Simple PDF generation using reportlab for text content
        try:
            # Read DOCX content (simplified)
            doc = Document(docx_path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # Create PDF
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            for text in content:
                para = Paragraph(text, styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 12))
            
            pdf_doc.build(story)
            return pdf_path
            
        except Exception as e:
            print(f"PDF generation error: {e}")
            return docx_path  # Return DOCX if PDF generation fails
    
    def _add_resume_styles(self, doc):
        """Add ATS-compliant styles to resume document"""
        styles = doc.styles
        
        # Modify existing styles for ATS compliance
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        
        heading1_style = styles['Heading 1']
        heading1_style.font.name = 'Calibri'
        heading1_style.font.size = Pt(16)
        heading1_style.font.bold = True
        
        heading2_style = styles['Heading 2']
        heading2_style.font.name = 'Calibri'
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True
        
        heading3_style = styles['Heading 3']
        heading3_style.font.name = 'Calibri'
        heading3_style.font.size = Pt(12)
        heading3_style.font.bold = True
    
    def _add_letter_styles(self, doc):
        """Add professional styles to cover letter document"""
        styles = doc.styles
        
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_after = Pt(6)
    
    def _format_experience_enhancements(self, enhancements: List[Dict]) -> str:
        """Format experience enhancements for display"""
        formatted = []
        for enhancement in enhancements:
            if isinstance(enhancement, dict):
                position = enhancement.get('position', 'Position')
                before = enhancement.get('before', '')
                after = enhancement.get('after', '')
                formatted.append(f"• {position}:")
                if before:
                    formatted.append(f"  Before: {before}")
                if after:
                    formatted.append(f"  After: {after}")
                formatted.append("")
        return '\n'.join(formatted)
    
    def cleanup_temp_files(self, older_than_hours: int = 24):
        """Clean up temporary files older than specified hours"""
        import time
        current_time = time.time()
        cutoff_time = current_time - (older_than_hours * 3600)
        
        for file_path in self.temp_dir.glob("*"):
            if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                try:
                    file_path.unlink()
                    print(f"Cleaned up temp file: {file_path}")
                except Exception as e:
                    print(f"Error cleaning up {file_path}: {e}")


# Helper functions for API usage
def generate_all_documents(analysis_results: Dict, user_id: str) -> Dict[str, str]:
    """Generate all professional documents from analysis results"""
    generator = DocumentGenerator()
    generated_files = {}
    
    try:
        # Generate resume DOCX
        if analysis_results.get('optimized_resume'):
            resume_file = generator.generate_resume_docx(
                analysis_results['optimized_resume'],
                f"resume_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            generated_files['resume_docx'] = resume_file
            generated_files['resume_pdf'] = generator.generate_pdf_from_docx(resume_file)
        
        # Generate cover letter DOCX
        if analysis_results.get('cover_letter'):
            cover_letter_file = generator.generate_cover_letter_docx(
                analysis_results['cover_letter'],
                f"cover_letter_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            generated_files['cover_letter_docx'] = cover_letter_file
            generated_files['cover_letter_pdf'] = generator.generate_pdf_from_docx(cover_letter_file)
        
        # Generate LinkedIn optimization DOCX
        if analysis_results.get('linkedin_optimization'):
            linkedin_file = generator.generate_linkedin_optimization_docx(
                analysis_results['linkedin_optimization'],
                f"linkedin_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            generated_files['linkedin_docx'] = linkedin_file
            generated_files['linkedin_pdf'] = generator.generate_pdf_from_docx(linkedin_file)
        
        # Generate interview preparation DOCX
        if analysis_results.get('interview_preparation'):
            interview_file = generator.generate_interview_prep_docx(
                analysis_results['interview_preparation'],
                f"interview_prep_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            generated_files['interview_docx'] = interview_file
            generated_files['interview_pdf'] = generator.generate_pdf_from_docx(interview_file)
        
        return generated_files
        
    except Exception as e:
        print(f"Error generating documents: {e}")
        return {}