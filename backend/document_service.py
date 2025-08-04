from docx import Document as DocxDocument
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import io
import os
from typing import Dict, Any, Optional
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self):
        self.upload_dir = os.getenv("UPLOAD_DIR", "uploads")
        os.makedirs(self.upload_dir, exist_ok=True)
    
    def generate_resume_docx(self, resume_data: Dict[str, Any], optimized_content: str = None) -> io.BytesIO:
        """Generate resume in DOCX format."""
        try:
            doc = DocxDocument()
            
            # Add personal information
            personal_info = resume_data.get('personal_info', {})
            name = personal_info.get('name', 'Your Name')
            
            # Title
            title = doc.add_heading(name, 0)
            title.alignment = 1  # Center alignment
            
            # Contact information
            contact_para = doc.add_paragraph()
            contact_info = []
            if personal_info.get('email'):
                contact_info.append(f"Email: {personal_info['email']}")
            if personal_info.get('phone'):
                contact_info.append(f"Phone: {personal_info['phone']}")
            if personal_info.get('location'):
                contact_info.append(f"Location: {personal_info['location']}")
            
            contact_para.add_run(' | '.join(contact_info))
            contact_para.alignment = 1  # Center alignment
            
            # Professional Summary
            if personal_info.get('summary'):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(personal_info['summary'])
            
            # Work Experience
            work_experience = resume_data.get('work_experience', [])
            if work_experience:
                doc.add_heading('Work Experience', level=1)
                for job in work_experience:
                    job_title = f"{job.get('title', '')} at {job.get('company', '')}"
                    if job.get('duration'):
                        job_title += f" ({job['duration']})"
                    doc.add_heading(job_title, level=2)
                    
                    if job.get('description'):
                        doc.add_paragraph(job['description'])
                    
                    if job.get('achievements'):
                        for achievement in job['achievements']:
                            p = doc.add_paragraph(achievement, style='ListBullet')
            
            # Education
            education = resume_data.get('education', [])
            if education:
                doc.add_heading('Education', level=1)
                for edu in education:
                    edu_title = f"{edu.get('degree', '')} - {edu.get('institution', '')}"
                    if edu.get('year'):
                        edu_title += f" ({edu['year']})"
                    doc.add_heading(edu_title, level=2)
                    
                    if edu.get('description'):
                        doc.add_paragraph(edu['description'])
            
            # Skills
            skills = resume_data.get('skills', [])
            if skills:
                doc.add_heading('Skills', level=1)
                skills_text = ', '.join(skills)
                doc.add_paragraph(skills_text)
            
            # Projects
            projects = resume_data.get('projects', [])
            if projects:
                doc.add_heading('Projects', level=1)
                for project in projects:
                    project_title = project.get('name', 'Project')
                    doc.add_heading(project_title, level=2)
                    
                    if project.get('description'):
                        doc.add_paragraph(project['description'])
                    
                    if project.get('technologies'):
                        tech_para = doc.add_paragraph()
                        tech_para.add_run('Technologies: ').bold = True
                        tech_para.add_run(', '.join(project['technologies']))
            
            # Save to BytesIO
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Error generating DOCX resume: {str(e)}")
            raise
    
    def generate_resume_pdf(self, resume_data: Dict[str, Any], optimized_content: str = None) -> io.BytesIO:
        """Generate resume in PDF format."""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                alignment=1,  # Center
                spaceAfter=12
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.darkblue,
                spaceAfter=6
            )
            
            # Personal Information
            personal_info = resume_data.get('personal_info', {})
            name = personal_info.get('name', 'Your Name')
            
            # Title
            story.append(Paragraph(name, title_style))
            
            # Contact information
            contact_info = []
            if personal_info.get('email'):
                contact_info.append(f"Email: {personal_info['email']}")
            if personal_info.get('phone'):
                contact_info.append(f"Phone: {personal_info['phone']}")
            if personal_info.get('location'):
                contact_info.append(f"Location: {personal_info['location']}")
            
            if contact_info:
                contact_text = ' | '.join(contact_info)
                story.append(Paragraph(contact_text, styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Professional Summary
            if personal_info.get('summary'):
                story.append(Paragraph('Professional Summary', heading_style))
                story.append(Paragraph(personal_info['summary'], styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Work Experience
            work_experience = resume_data.get('work_experience', [])
            if work_experience:
                story.append(Paragraph('Work Experience', heading_style))
                for job in work_experience:
                    job_title = f"{job.get('title', '')} at {job.get('company', '')}"
                    if job.get('duration'):
                        job_title += f" ({job['duration']})"
                    
                    story.append(Paragraph(job_title, styles['Heading3']))
                    
                    if job.get('description'):
                        story.append(Paragraph(job['description'], styles['Normal']))
                    
                    if job.get('achievements'):
                        for achievement in job['achievements']:
                            story.append(Paragraph(f"• {achievement}", styles['Normal']))
                    
                    story.append(Spacer(1, 6))
            
            # Education
            education = resume_data.get('education', [])
            if education:
                story.append(Paragraph('Education', heading_style))
                for edu in education:
                    edu_title = f"{edu.get('degree', '')} - {edu.get('institution', '')}"
                    if edu.get('year'):
                        edu_title += f" ({edu['year']})"
                    
                    story.append(Paragraph(edu_title, styles['Heading3']))
                    if edu.get('description'):
                        story.append(Paragraph(edu['description'], styles['Normal']))
                
                story.append(Spacer(1, 12))
            
            # Skills
            skills = resume_data.get('skills', [])
            if skills:
                story.append(Paragraph('Skills', heading_style))
                skills_text = ', '.join(skills)
                story.append(Paragraph(skills_text, styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Projects
            projects = resume_data.get('projects', [])
            if projects:
                story.append(Paragraph('Projects', heading_style))
                for project in projects:
                    project_title = project.get('name', 'Project')
                    story.append(Paragraph(project_title, styles['Heading3']))
                    
                    if project.get('description'):
                        story.append(Paragraph(project['description'], styles['Normal']))
                    
                    if project.get('technologies'):
                        tech_text = f"Technologies: {', '.join(project['technologies'])}"
                        story.append(Paragraph(tech_text, styles['Normal']))
                    
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Error generating PDF resume: {str(e)}")
            raise
    
    def generate_cover_letter_docx(self, cover_letter_data: Dict[str, Any]) -> io.BytesIO:
        """Generate cover letter in DOCX format."""
        try:
            doc = DocxDocument()
            
            # Header with date and contact info
            header = doc.add_paragraph()
            header.add_run(datetime.now().strftime("%B %d, %Y")).bold = True
            header.alignment = 2  # Right alignment
            
            doc.add_paragraph()  # Empty line
            
            # Employer information
            if cover_letter_data.get('company_name'):
                doc.add_paragraph(f"Hiring Manager\n{cover_letter_data['company_name']}")
            
            doc.add_paragraph()  # Empty line
            
            # Salutation
            salutation = cover_letter_data.get('salutation', 'Dear Hiring Manager,')
            doc.add_paragraph(salutation)
            
            # Cover letter content
            content = cover_letter_data.get('content', '')
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Closing
            doc.add_paragraph()
            closing = cover_letter_data.get('closing', 'Sincerely,')
            doc.add_paragraph(closing)
            
            # Signature
            signature = cover_letter_data.get('signature', 'Your Name')
            doc.add_paragraph(signature)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Error generating DOCX cover letter: {str(e)}")
            raise
    
    def generate_cover_letter_pdf(self, cover_letter_data: Dict[str, Any]) -> io.BytesIO:
        """Generate cover letter in PDF format."""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Header with date
            story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Employer information
            if cover_letter_data.get('company_name'):
                story.append(Paragraph(f"Hiring Manager<br/>{cover_letter_data['company_name']}", styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Salutation
            salutation = cover_letter_data.get('salutation', 'Dear Hiring Manager,')
            story.append(Paragraph(salutation, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Cover letter content
            content = cover_letter_data.get('content', '')
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['Normal']))
                    story.append(Spacer(1, 12))
            
            # Closing
            closing = cover_letter_data.get('closing', 'Sincerely,')
            story.append(Paragraph(closing, styles['Normal']))
            story.append(Spacer(1, 24))
            
            # Signature
            signature = cover_letter_data.get('signature', 'Your Name')
            story.append(Paragraph(signature, styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Error generating PDF cover letter: {str(e)}")
            raise

# Create global instance
document_generator = DocumentGenerator()